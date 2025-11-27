import os
import sys
import io
import json
import time
import math
import zipfile
import tempfile
import logging
from logging.handlers import RotatingFileHandler
import threading
import subprocess
import platform
import uuid
from datetime import datetime
import shutil
from concurrent.futures import ThreadPoolExecutor, as_completed

from dotenv import load_dotenv
from flask import (
    Flask, render_template, request, redirect, flash, url_for,
    jsonify, send_file, after_this_request, g
)
from flask_mail import Mail, Message
from werkzeug.utils import secure_filename

import fitz  # PyMuPDF
from PIL import Image
from fpdf import FPDF
from pdf2docx import Converter
import pdfplumber
import pytesseract
import camelot
import pandas as pd

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "chave-padrao-fallback")
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024

# ---------------------------- CONFIG ----------------------------
app.config['MAIL_SERVER'] = os.getenv("MAIL_SERVER")
app.config['MAIL_PORT'] = int(os.getenv("MAIL_PORT", 587))
app.config['MAIL_USE_TLS'] = os.getenv("MAIL_USE_TLS", "True").lower() in ("true", "1", "yes")
app.config['MAIL_USERNAME'] = os.getenv("MAIL_USERNAME")
app.config['MAIL_PASSWORD'] = os.getenv("MAIL_PASSWORD")
app.config['MAIL_DEFAULT_SENDER'] = (os.getenv("MAIL_DEFAULT_SENDER_NAME"), os.getenv("MAIL_DEFAULT_SENDER_EMAIL"))
mail = Mail(app)

# =================== LOGGING GLOBAL ===================

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
LOG_DIR = os.path.join(PROJECT_DIR, "logs")
os.makedirs(LOG_DIR, exist_ok=True)

LOG_FILE = os.path.join(LOG_DIR, "zg_pdf.log")

# Formato bonitinho
formatter = logging.Formatter(
    '[%(asctime)s] [%(levelname)s] [%(name)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

# 🔁 Rotação por tamanho (ex: 10 MB, guarda 5 arquivos antigos)
file_handler = RotatingFileHandler(
    LOG_FILE,
    maxBytes=10 * 1024 * 1024,  # 10 MB
    backupCount=5,              # zg_pdf.log, zg_pdf.log.1, ... .5
    encoding="utf-8"
)
file_handler.setFormatter(formatter)
file_handler.setLevel(logging.INFO)

# Log no console (vai pro journalctl)
console_handler = logging.StreamHandler()
console_handler.setFormatter(formatter)
console_handler.setLevel(logging.INFO)

# Root logger
root_logger = logging.getLogger()
root_logger.setLevel(logging.INFO)

# Evita duplicar handler se recarregar app em dev
if not any(isinstance(h, RotatingFileHandler) for h in root_logger.handlers):
    root_logger.addHandler(file_handler)
if not any(isinstance(h, logging.StreamHandler) for h in root_logger.handlers):
    root_logger.addHandler(console_handler)

# Logger “principal” da app (opcional, se quiser usar logging.getLogger("zg_pdf"))
log = logging.getLogger("zg_pdf")

# Diretórios
WORK_DIR = tempfile.gettempdir() if getattr(sys, 'frozen', False) else os.getcwd()
UPLOAD_FOLDER = os.path.join(WORK_DIR, 'uploads')
PROCESSED_FOLDER = os.path.join(WORK_DIR, 'processed')

for folder in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
    os.makedirs(folder, exist_ok=True)

# Logs / tasks
COMPRESSION_LOG_FILE = 'compression_log.json'
log_lock = threading.Lock()
tasks_lock = threading.Lock()
tasks = {}

# Ajuste Ghostscript por SO
GS_CMD = "gswin64c" if platform.system() == "Windows" else "gs"

# Pytesseract: sem caminho fixo no Linux
if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


# ---------------------------- Utils ----------------------------
def save_compression_log(log_data: dict) -> None:
    """Append seguro em JSON (cria se não existir)."""
    with log_lock:
        if not os.path.exists(COMPRESSION_LOG_FILE):
            with open(COMPRESSION_LOG_FILE, 'w', encoding='utf-8') as f:
                json.dump([], f, ensure_ascii=False)
        with open(COMPRESSION_LOG_FILE, 'r+', encoding='utf-8') as f:
            data = json.load(f)
            data.append(log_data)
            f.seek(0)
            json.dump(data, f, indent=4, ensure_ascii=False)
            f.truncate()


def get_pdf_page_count(pdf_path: str) -> int:
    """Conta páginas com pypdf (robusto)."""
    try:
        from pypdf import PdfReader
        return len(PdfReader(pdf_path).pages)
    except Exception as e:
        logging.error(f"Erro ao obter número de páginas do PDF {pdf_path}: {e}")
        return 0


def get_estimated_compression_ratio(compression_type: str) -> float:
    ratios = {
        'screen': 0.20,  # ~80% de compressão
        'ebook': 0.40,   # ~60% de compressão
    }
    return ratios.get(compression_type, 0.30)


# ---------------------------- Templates ----------------------------
@app.route('/')
def index():
    return render_template('main.html')

@app.route('/compress', methods=['GET'])
def compress_form():
    return render_template('compress.html')

@app.route('/split', methods=['GET'])
def split_form():
    return render_template('split.html')

@app.route('/convert', methods=['GET'])
def convert_form():
    return render_template('convert.html')

@app.route('/merge', methods=['GET'])
def merge_form():
    return render_template('merge.html')

@app.route('/organize', methods=['GET'])
def organize_form():
    return render_template('organize.html')

@app.route('/extract', methods=['GET'])
def extract_form():
    return render_template('extract.html')


# ---------------------------- Feedback ----------------------------
@app.route('/enviar-feedback', methods=['POST'])
def enviar_feedback():
    nome = request.form.get('name')
    email = request.form.get('email')
    setor = request.form.get('sector')
    mensagem = request.form.get('message')

    corpo_email_html = f"""
    <html>
    <body style="font-family: Arial, Helvetica, sans-serif; color: #444; background: #f9fafc; padding: 20px; margin: 0;">
      <div style="max-width: 600px; margin: auto; background: #fff; border-radius: 12px; box-shadow: 0 4px 15px rgba(0,0,0,0.1); padding: 30px;">
        <h2 style="-webkit-background-clip:text;-webkit-text-fill-color:transparent;background:linear-gradient(90deg,#0052cc,#007bff);font-weight:700;font-size:28px;margin-bottom:25px;display:flex;align-items:center;gap:10px;">📬 Novo feedback recebido</h2>
        <p><strong>👤 Nome:</strong> {nome}</p>
        <p><strong>📧 E-mail:</strong> {email}</p>
        <p><strong>🏢 Setor:</strong> {setor}</p>
        <p><strong>💬 Mensagem:</strong></p>
        <p style="background:#f5f7fa;border-left:5px solid #007bff;padding:15px 20px;border-radius:8px;white-space:pre-wrap;color:#333;">{mensagem}</p>
        <hr style="border:none;border-top:1px solid #ddd;margin:30px 0;">
        <p style="font-size:12px;color:#999;text-align:center;font-style:italic;">Enviado automaticamente pelo sistema de feedback do site.</p>
      </div>
    </body>
    </html>
    """

    try:
        msg = Message(subject="Feedback do Site", recipients=["ti@zavagnagralha.com.br"])
        msg.body = "Você recebeu um novo feedback!"
        msg.html = corpo_email_html
        mail.send(msg)
        flash("Mensagem enviada com sucesso! Obrigado pelo seu feedback.", "success")
        return '', 200
    except Exception as e:
        flash(f"Erro ao enviar mensagem: {e}", "error")
        return str(e), 500

# ---------------------------- Compressão ----------------------------
@app.route("/compress", methods=["POST"])
def compress():
    # Tenta pegar vários arquivos (nova forma)
    files = request.files.getlist('files')

    # Retrocompat: se vier só 'file'
    if not files:
        single = request.files.get('file')
        if single:
            files = [single]
        else:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400

    MAX_FILES = 10
    MAX_TOTAL_MB = 1024  # MB

    if len(files) > MAX_FILES:
        return jsonify({
            'error': f'Você enviou {len(files)} arquivos. O limite é {MAX_FILES}.'
        }), 400

    task_id = str(uuid.uuid4())

    jobs = []
    total_bytes = 0

    for f in files:
        filename = secure_filename(f.filename)
        input_path = os.path.join(UPLOAD_FOLDER, f"{task_id}_{filename}")
        output_path = os.path.join(PROCESSED_FOLDER, f"comprimidoZG_{filename}")
        f.save(input_path)

        size = os.path.getsize(input_path)
        total_bytes += size

        jobs.append({
            "filename": filename,
            "input_path": input_path,
            "output_path": output_path,
            "input_size": size,
        })

    total_mb = total_bytes / (1024 * 1024)
    if total_mb > MAX_TOTAL_MB:
        for job in jobs:
            try:
                os.remove(job["input_path"])
            except Exception:
                pass
        return jsonify({
            'error': f'O total enviado é {total_mb:.2f} MB, maior que o limite de {MAX_TOTAL_MB} MB.'
        }), 400

    compression_type = request.form.get('compression', 'screen')

    with tasks_lock:
        tasks[task_id] = {
            'percent': 0,
            'status': 'Iniciando compressão...',
            'file': None,
            'error': None,
            'summary': None,
        }

    thread = threading.Thread(
        target=compress_task_thread_many,
        args=(task_id, jobs, compression_type),
        daemon=True,
    )
    thread.start()

    return jsonify({'task_id': task_id})


@app.route("/progress/<task_id>")
def progress(task_id):
    with tasks_lock:
        task = tasks.get(task_id)
    if task:
        if task.get('error'):
            return jsonify(task), 500
        return jsonify(task)
    logging.warning(f"Tarefa {task_id} não encontrada para progresso.")
    return jsonify({'error': 'Tarefa não encontrada'}), 404


@app.route("/download/<task_id>")
def download(task_id):
    logging.info(f"Requisição de download recebida para task_id: {task_id}")
    with tasks_lock:
        task = tasks.get(task_id)

    if not task:
        logging.warning(f"Download repetido ou tarefa expirada: {task_id}")
        # 410 = Gone (já existiu um dia, mas não mais)
        return jsonify({'error': 'Download já realizado ou tarefa expirada.'}), 410

    file_path = task.get('file')
    if not file_path or not os.path.exists(file_path):
        logging.error(f"Download: Arquivo '{file_path}' não encontrado.")
        return jsonify({'error': 'Arquivo não disponível'}), 404

    logging.info(f"Servindo arquivo: {file_path}")

    @after_this_request
    def cleanup(response):
        try:
            os.remove(file_path)
            logging.info(f"Arquivo {file_path} removido.")
        except Exception as e:
            logging.warning(f"Erro ao remover arquivo {file_path}: {e}")
        with tasks_lock:
            tasks.pop(task_id, None)
        return response

    return send_file(file_path, as_attachment=True)

def compress_single_pdf(job, compression_type):
    """
    job: {
      "filename": str,
      "input_path": str,
      "output_path": str,
      "input_size": int (bytes)
    }
    Retorna um dict com métricas e o path de saída.
    """
    input_path = job["input_path"]
    output_path = job["output_path"]
    filename = job["filename"]

    logging.info(f"Compressão (single) iniciada para: {filename}")

    initial_file_size = os.path.getsize(input_path)
    compression_ratio = get_estimated_compression_ratio(compression_type)
    estimated_final_size = max(int(initial_file_size * compression_ratio), 1)

    def run_gs(current_input_path, cmd_output_path, compression_type_override, extra_flags=None):
        gs_path = shutil.which("gs") or "/usr/bin/gs"
        if not os.path.exists(gs_path):
            raise FileNotFoundError(
                f"Ghostscript não encontrado em PATH. "
                f"Tente instalar ou definir o caminho. Procurado: {gs_path}"
            )

        command = [
            gs_path,
            "-sDEVICE=pdfwrite",
            "-dCompatibilityLevel=1.4",
            f"-dPDFSETTINGS=/{compression_type_override}",
            "-dNOPAUSE",
            "-dBATCH",
            "-dQUIET",
            f"-sOutputFile={cmd_output_path}",
            current_input_path,
        ]
        if extra_flags:
            command += extra_flags

        return subprocess.Popen(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )

    # Chamada principal
    process = run_gs(input_path, output_path, compression_type)

    stdout, stderr = process.communicate()
    if process.returncode != 0:
        raise subprocess.CalledProcessError(
            process.returncode,
            process.args,
            output=stdout,
            stderr=stderr,
        )

    if not os.path.exists(output_path):
        raise FileNotFoundError(f"Arquivo de saída não foi gerado: {output_path}")

    # --------- fallback agressivo se não tiver ganho ---------
    final_file_size = os.path.getsize(output_path)
    if final_file_size >= initial_file_size:
        logging.warning(f"Compressão ineficaz em {filename}. Tentando fallback agressivo.")
        fallback_path = output_path.replace('.pdf', '_fallback.pdf')
        aggressive_flags = [
            '-dDownsampleColorImages=true', '-dColorImageResolution=72', '-dAutoFilterColorImages=false', '-dColorImageFilter=/DCTEncode',
            '-dGrayImageResolution=72', '-dDownsampleGrayImages=true', '-dAutoFilterGrayImages=false', '-dGrayImageFilter=/DCTEncode',
            '-dMonoImageResolution=72', '-dDownsampleMonoImages=true',
        ]
        fallback_process = run_gs(input_path, fallback_path, "screen", aggressive_flags)
        fb_stdout, fb_stderr = fallback_process.communicate()

        if os.path.exists(fallback_path):
            fallback_size = os.path.getsize(fallback_path)
            if fallback_size < initial_file_size:
                os.replace(fallback_path, output_path)
                final_file_size = fallback_size
            else:
                os.replace(input_path, output_path)
                final_file_size = initial_file_size
                try:
                    os.remove(fallback_path)
                except Exception:
                    pass
        else:
            os.replace(input_path, output_path)
            final_file_size = initial_file_size

    # --------- métricas finais ---------
    size_reduction_bytes = initial_file_size - final_file_size
    size_reduction_mb = round(size_reduction_bytes / (1024 * 1024), 2)
    input_mb = round(initial_file_size / (1024 * 1024), 2)
    output_mb = round(final_file_size / (1024 * 1024), 2)
    if initial_file_size > 0:
        size_reduction_pct = round((size_reduction_bytes / initial_file_size) * 100, 2)
    else:
        size_reduction_pct = 0.0

    pages = get_pdf_page_count(output_path)

    # limpa input individual
    try:
        os.remove(input_path)
    except Exception as e:
        logging.warning(f"Não foi possível remover o arquivo de entrada {input_path}: {e}")

    return {
        "filename": filename,
        "output_path": output_path,
        "input_mb": input_mb,
        "output_mb": output_mb,
        "reduction_mb": size_reduction_mb,
        "reduction_pct": size_reduction_pct,
        "pages": pages,
    }

def compress_task_thread_many(current_task_id, jobs, compression_type):
    start_time_global = time.time()
    started_at = datetime.utcnow().isoformat() + "Z"
    logging.info(f"Thread de compressão iniciada para a tarefa: {current_task_id} ({len(jobs)} arquivos)")

    all_logs = []
    total_input_mb = 0.0
    total_output_mb = 0.0

    try:
        total_jobs = len(jobs)
        MAX_PARALLEL = 3  # ajusta conforme o poder da máquina

        with tasks_lock:
            if current_task_id in tasks:
                tasks[current_task_id]['status'] = f"Iniciando compressão de {total_jobs} arquivos..."
                tasks[current_task_id]['percent'] = 0

        with ThreadPoolExecutor(max_workers=MAX_PARALLEL) as executor:
            future_to_job = {
                executor.submit(compress_single_pdf, job, compression_type): job
                for job in jobs
            }

            done_count = 0

            for future in as_completed(future_to_job):
                job = future_to_job[future]
                filename = job["filename"]

                try:
                    result = future.result()
                    all_logs.append(result)

                    total_input_mb += result["input_mb"]
                    total_output_mb += result["output_mb"]

                    done_count += 1
                    percent = int((done_count / total_jobs) * 100)

                    with tasks_lock:
                        if current_task_id in tasks:
                            tasks[current_task_id]['percent'] = min(percent, 99)
                            tasks[current_task_id]['status'] = (
                                f"Comprimindo arquivos... ({done_count}/{total_jobs})"
                            )

                except Exception as e:
                    logging.error(f"Erro ao comprimir arquivo {filename} na tarefa {current_task_id}: {e}")
                    done_count += 1
                    with tasks_lock:
                        if current_task_id in tasks:
                            tasks[current_task_id]['status'] = f"Erro ao comprimir \"{filename}\". Prosseguindo com os demais..."
                            tasks[current_task_id]['percent'] = int((done_count / total_jobs) * 100)

        # terminou todos: gera ZIP
        zip_path = os.path.join(PROCESSED_FOLDER, f"comprimidosZG_{current_task_id}.zip")
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for res in all_logs:
                out_path = res["output_path"]
                if os.path.exists(out_path):
                    arcname = os.path.basename(out_path)
                    zipf.write(out_path, arcname=arcname)
                    try:
                        os.remove(out_path)
                    except Exception:
                        pass

        end_time_global = time.time()
        time_taken_global = round(end_time_global - start_time_global, 2)

        if total_input_mb > 0:
            total_reduction_pct = round((total_input_mb - total_output_mb) / total_input_mb * 100, 2)
        else:
            total_reduction_pct = 0.0

        with tasks_lock:
            if current_task_id in tasks:
                tasks[current_task_id]['percent'] = 100
                tasks[current_task_id]['status'] = 'Compressão concluída! Preparando download...'
                tasks[current_task_id]['file'] = zip_path
                tasks[current_task_id]['summary'] = {
                    'files_count': len(all_logs),
                    'input_mb': total_input_mb,
                    'output_mb': total_output_mb,
                    'reduction_pct': total_reduction_pct,
                    'time_s': time_taken_global,
                }

        # logs unitários
        for res in all_logs:
            log_data = {
                'task_id': current_task_id,
                'input_file': res["filename"],
                'output_file': os.path.basename(res["output_path"]),
                'compression_type': compression_type,
                'time_taken_seconds': time_taken_global,  # ou tira se não fizer questão
                'pages': res["pages"],
                'input_file_size_mb': res["input_mb"],
                'output_file_size_mb': res["output_mb"],
                'size_reduction_mb': res["reduction_mb"],
                'size_reduction_pct': res["reduction_pct"],
                'status': 'Concluído!',
                'started_at': started_at,
                'finished_at': datetime.utcnow().isoformat() + "Z",
                'batch_files_count': len(all_logs),       # 👈 novo
                'is_batch': True,                         # 👈 novo
            }
            save_compression_log(log_data)

    except Exception as e:
        logging.error(f"Erro inesperado na tarefa {current_task_id}: {e}")
        with tasks_lock:
            if current_task_id in tasks:
                tasks[current_task_id]['status'] = 'Erro interno'
                tasks[current_task_id]['error'] = str(e)
                tasks[current_task_id]['percent'] = -1
    finally:
        # garante remoção de inputs que por acaso sobraram
        for job in jobs:
            if os.path.exists(job["input_path"]):
                try:
                    os.remove(job["input_path"])
                except Exception as e:
                    logging.warning(f"Não foi possível remover o arquivo de entrada {job['input_path']}: {e}")

# ---------------------------- Conversão ----------------------------
def extract_text_from_pdf(input_path, lang='por'):
    text = ""
    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            extracted_text = page.extract_text()
            if extracted_text and extracted_text.strip():
                text += extracted_text + "\n"
            else:
                # OCR fallback
                pil_image = page.to_image().original.convert('RGB')
                ocr_text = pytesseract.image_to_string(pil_image, lang=lang)
                text += ocr_text + "\n"
    return text.strip()


def is_pdf_scanned(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        total_text = ""
        for page in pdf.pages:
            t = page.extract_text()
            if t and t.strip():
                total_text += t.strip()
        return len(total_text) < 200


def pdf_to_docx(input_pdf_path, output_docx_path, lang='por', conf_threshold=50):
    # 1) PDF -> DOCX (layout)
    docx_buffer = io.BytesIO()
    cv = Converter(input_pdf_path)
    cv.convert(docx_buffer, start=0, end=None)
    cv.close()
    docx_buffer.seek(0)

    from docx import Document
    word_doc = Document(docx_buffer)

    # 2) OCR por página se necessário
    pdf_doc = fitz.open(input_pdf_path)

    def is_image_textual(img, lang=lang, conf_threshold=conf_threshold):
        data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)
        confs = []
        for c in data['conf']:
            try:
                v = float(c)
                if v > 0:
                    confs.append(v)
            except Exception:
                pass
        if confs:
            return (sum(confs) / len(confs)) >= conf_threshold
        return False

    for i, page in enumerate(pdf_doc):
        text_in_page = page.get_text().strip()
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        if not text_in_page and is_image_textual(img, lang=lang):
            ocr_text = pytesseract.image_to_string(img, lang=lang).strip()
            if i < len(word_doc.paragraphs):
                p = word_doc.paragraphs[i]
                # Limpeza simples: remove runs e substitui
                for r in p.runs:
                    r.text = ""
                p.add_run(ocr_text)
            else:
                word_doc.add_paragraph(ocr_text)

    word_doc.save(output_docx_path)


conversion_map = {
    'pdf': ['docx', 'xlsx', 'txt'],
    'docx': ['pdf', 'txt'],
    'xlsx': ['pdf'],
    'pptx': ['pdf'],
    'jpg': ['pdf'],
    'jpeg': ['pdf'],
    'png': ['pdf'],
}


@app.route('/convert-extensions', methods=['GET'])
def get_supported_extensions():
    return jsonify({'extensions': list(conversion_map.keys())})


@app.route('/upload-conversion', methods=['POST'])
def upload_conversion():
    file = request.files.get('file')
    if not file:
        return jsonify({'error': 'Nenhum arquivo enviado.'}), 400

    filename = secure_filename(file.filename)
    ext = os.path.splitext(filename)[1].lower().replace('.', '')
    available_options = conversion_map.get(ext)

    if not available_options:
        return jsonify({
            'filename': filename,
            'extension': ext,
            'message': 'Tipo de arquivo não suportado para conversão',
            'options': []
        }), 200

    task_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_FOLDER, f"{task_id}_{filename}")
    file.save(input_path)

    is_scanned = is_pdf_scanned(input_path) if ext == 'pdf' else False

    return jsonify({
        'filename': filename,
        'extension': ext,
        'task_id': task_id,
        'message': 'Arquivo recebido com sucesso!',
        'options': available_options,
        'is_scanned': is_scanned
    }), 200


def _find_soffice():
    candidates = [
        shutil.which("soffice"),
        "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
        shutil.which("libreoffice"),   # snap costuma expor esse
        "/snap/bin/libreoffice",
    ]
    for p in candidates:
        if p and os.path.exists(p):
            return p
    return None

def convert_office_to_pdf_libreoffice(input_path):
    """Office -> PDF via LibreOffice (Linux)."""
    soffice = _find_soffice()
    if not soffice:
        raise FileNotFoundError(
            "LibreOffice não encontrado. Instale 'libreoffice' e/ou ajuste PATH. "
            "Procurei por: soffice/libreoffice em /usr/bin, /usr/lib/libreoffice, /snap/bin."
        )
    # Usa 'libreoffice' se for o que achou
    cmd = [soffice, '--headless', '--convert-to', 'pdf', '--outdir',
           os.path.dirname(input_path), input_path]
    try:
        res = subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    except FileNotFoundError as e:
        raise FileNotFoundError(f"Executável não encontrado: {soffice}") from e
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Falha na conversão LO. stdout:\n{e.stdout}\nstderr:\n{e.stderr}") from e
    return os.path.splitext(input_path)[0] + '.pdf'


def convert_office_to_pdf(input_path):
    """Despacha por SO. No Linux usa LibreOffice, no Windows tu pode plugar automação Office se quiser (removido aqui)."""
    if platform.system() == 'Windows':
        raise RuntimeError("Conversão Office->PDF no Windows não está habilitada nesta build.")
    return convert_office_to_pdf_libreoffice(input_path)

def run_conversion(input_path: str, target_format: str) -> tuple[bytes, str]:
    """
    Converte um único arquivo salvo em disco.
    Retorna (conteúdo_em_bytes, nome_sugerido_para_download).
    NÃO remove o arquivo original (isso é responsabilidade da rota).
    """
    ext = os.path.splitext(input_path)[1].lower().replace('.', '')
    temp_files_to_delete = []
    output_buffer = io.BytesIO()
    download_name = f"convertido.{target_format}"

    try:
        if ext == 'pdf':
            text = extract_text_from_pdf(input_path)

            if target_format == 'docx':
                temp_docx_path = os.path.join(PROCESSED_FOLDER, f"{uuid.uuid4()}.docx")
                pdf_to_docx(input_path, temp_docx_path)
                temp_files_to_delete.append(temp_docx_path)
                with open(temp_docx_path, 'rb') as f:
                    output_buffer.write(f.read())
                download_name = os.path.basename(temp_docx_path)

            elif target_format == 'xlsx':
                lines = text.strip().split("\n")
                df_list = []
                tables = camelot.read_pdf(input_path, pages='all', flavor='lattice')
                if tables.n == 0:
                    tables = camelot.read_pdf(input_path, pages='all', flavor='stream')
                if tables.n == 0:
                    df_list.append(pd.DataFrame({'Conteúdo': lines}))
                else:
                    for tbl in tables:
                        df = tbl.df.copy()
                        df = df.applymap(lambda x: ' '.join(str(x).split()))
                        df_list.append(df)
                with pd.ExcelWriter(output_buffer, engine='xlsxwriter') as writer:
                    for idx, df in enumerate(df_list):
                        df.to_excel(writer, sheet_name=f'Tabela_{idx+1}'[:31], index=False)
                    pd.DataFrame({'Texto extraído': lines}).to_excel(writer, sheet_name='Texto_RAW', index=False)
                download_name = "convertido.xlsx"

            elif target_format == 'txt':
                output_buffer.write(text.encode('utf-8'))
                download_name = "convertido.txt"

            else:
                raise ValueError('Conversão não suportada para PDF.')

        elif ext in ['docx', 'xlsx', 'pptx']:
            if target_format == 'pdf':
                pdf_path = convert_office_to_pdf(input_path)
                temp_files_to_delete.append(pdf_path)
                with open(pdf_path, 'rb') as f:
                    output_buffer.write(f.read())
                download_name = os.path.basename(pdf_path)
            elif target_format == 'txt' and ext == 'docx':
                import docx
                doc = docx.Document(input_path)
                full_text = '\n'.join([p.text for p in doc.paragraphs])
                output_buffer.write(full_text.encode('utf-8'))
                download_name = "convertido.txt"
            elif target_format == 'image':
                raise RuntimeError('Conversão Office → Imagem não implementada.')
            else:
                raise ValueError('Conversão não suportada.')

        elif ext in ['jpg', 'jpeg', 'png']:
            if target_format == 'pdf':
                image = Image.open(input_path).convert('RGB')
                img_w_px, img_h_px = image.size

                if img_w_px >= img_h_px:
                    orientation = 'L'; page_w_mm, page_h_mm = 297, 210
                else:
                    orientation = 'P'; page_w_mm, page_h_mm = 210, 297

                margin_mm = 10
                max_w_mm = page_w_mm - 2 * margin_mm
                max_h_mm = page_h_mm - 2 * margin_mm

                scale_w = max_w_mm / img_w_px
                scale_h = max_h_mm / img_h_px
                scale = min(scale_w, scale_h)

                disp_w_mm = img_w_px * scale
                disp_h_mm = img_h_px * scale
                x_mm = (page_w_mm - disp_w_mm) / 2.0
                y_mm = (page_h_mm - disp_h_mm) / 2.0

                pdf = FPDF(orientation=orientation, unit='mm', format='A4')
                pdf.set_auto_page_break(False)
                pdf.add_page()

                with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp:
                    temp_path = tmp.name
                image.save(temp_path, 'JPEG', quality=95)

                pdf.image(temp_path, x=x_mm, y=y_mm, w=disp_w_mm, h=disp_h_mm)

                try:
                    os.remove(temp_path)
                except Exception:
                    pass

                pdf_data = pdf.output(dest='S')
                if isinstance(pdf_data, str):
                    pdf_data = pdf_data.encode('latin1')

                output_buffer.write(pdf_data)
                download_name = "convertido.pdf"

            else:
                raise ValueError('Conversão não suportada para imagem.')

        else:
            raise ValueError('Formato de entrada não reconhecido.')

        return output_buffer.getvalue(), download_name

    finally:
        for temp_file in temp_files_to_delete:
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except Exception as e:
                    logging.warning(f"Erro removendo arquivo temporário: {temp_file} - {e}")


@app.route('/execute-conversion', methods=['POST'])
def execute_conversion():
    data = request.json or {}
    task_id = data.get('task_id')
    target_format = data.get('target_format')

    if not task_id or not target_format:
        return jsonify({'error': 'Parâmetros inválidos para conversão.'}), 400

    file_match = None
    for f in os.listdir(UPLOAD_FOLDER):
        if f.startswith(task_id):
            file_match = f
            break
    if not file_match:
        return jsonify({'error': 'Arquivo temporário não encontrado.'}), 404

    input_path = os.path.join(UPLOAD_FOLDER, file_match)

    try:
        content, download_name = run_conversion(input_path, target_format)
        # limpa upload temporário
        try:
            if os.path.exists(input_path):
                os.remove(input_path)
        except Exception as e:
            logging.warning(f"Erro removendo arquivo original: {input_path} - {e}")

        return send_file(
            io.BytesIO(content),
            as_attachment=True,
            download_name=download_name
        )

    except ValueError as ve:
        logging.warning(f"Conversão não suportada: {ve}")
        return jsonify({'error': str(ve)}), 400
    except Exception as e:
        logging.error(f"Erro na conversão: {e}")
        return jsonify({'error': 'Erro interno na conversão.'}), 500

@app.route('/upload-conversion-batch', methods=['POST'])
def upload_conversion_batch():
    """
    Recebe vários arquivos e devolve:
    - batch_id
    - items: [{ filename, extension, supported, task_id, options, is_scanned }]
    """
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'Nenhum arquivo enviado.'}), 400

    batch_id = str(uuid.uuid4())
    items = []

    for file in files:
        filename = secure_filename(file.filename)
        ext = os.path.splitext(filename)[1].lower().replace('.', '')
        available_options = conversion_map.get(ext)

        # não suportado
        if not available_options:
            items.append({
                'filename': filename,
                'extension': ext,
                'supported': False,
                'options': [],
                'is_scanned': False,
            })
            continue

        task_id = str(uuid.uuid4())
        stored_name = f"{batch_id}_{task_id}_{filename}"
        input_path = os.path.join(UPLOAD_FOLDER, stored_name)
        file.save(input_path)

        is_scanned = is_pdf_scanned(input_path) if ext == 'pdf' else False

        items.append({
            'filename': filename,
            'extension': ext,
            'supported': True,
            'task_id': task_id,
            'options': available_options,
            'is_scanned': is_scanned,
        })

    return jsonify({
        'batch_id': batch_id,
        'items': items
    }), 200

@app.route('/execute-conversion-batch', methods=['POST'])
def execute_conversion_batch():
    """
    Executa conversão em lote.
    Espera JSON:
    {
      "batch_id": "....",
      "targets": [
        {"task_id": "...", "target_format": "pdf"},
        ...
      ]
    }
    Retorna um ZIP com os arquivos convertidos.
    """
    data = request.json or {}
    batch_id = data.get('batch_id')
    targets = data.get('targets') or []

    if not batch_id or not isinstance(targets, list) or not targets:
        return jsonify({'error': 'batch_id e targets são obrigatórios.'}), 400

    zip_buffer = io.BytesIO()
    results = []

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as z:
        for t in targets:
            task_id = t.get('task_id')
            target_format = t.get('target_format')

            if not task_id or not target_format:
                continue

            prefix = f"{batch_id}_{task_id}_"
            file_match = None
            for f in os.listdir(UPLOAD_FOLDER):
                if f.startswith(prefix):
                    file_match = f
                    break

            if not file_match:
                logging.warning(f"Arquivo temporário não encontrado para task_id={task_id}")
                results.append({'task_id': task_id, 'status': 'missing'})
                continue

            input_path = os.path.join(UPLOAD_FOLDER, file_match)

            try:
                content, _download_name = run_conversion(input_path, target_format)

                # Recupera nome original: batchid_taskid_nomeoriginal.ext
                try:
                    original_name = file_match.split('_', 2)[2]
                except IndexError:
                    original_name = os.path.basename(file_match)

                base, _ = os.path.splitext(original_name)
                out_name = f"{base}.{target_format}"

                z.writestr(out_name, content)
                results.append({'task_id': task_id, 'status': 'ok', 'output': out_name})

            except ValueError as ve:
                logging.warning(f"Conversão não suportada em lote (task_id={task_id}): {ve}")
                results.append({'task_id': task_id, 'status': 'error', 'error': str(ve)})
            except Exception as e:
                logging.error(f"Erro convertendo em lote (task_id={task_id}): {e}")
                results.append({'task_id': task_id, 'status': 'error', 'error': 'Erro interno'})
            finally:
                # Remove o upload original
                try:
                    if os.path.exists(input_path):
                        os.remove(input_path)
                except Exception as e:
                    logging.warning(f"Erro removendo arquivo original em lote: {input_path} - {e}")

    zip_buffer.seek(0)
    zip_name = f"conversao_lote_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"

    return send_file(
        zip_buffer,
        as_attachment=True,
        download_name=zip_name,
        mimetype='application/zip'
    )


# ---------------------------- Merge / Split / Organize ----------------------------
@app.route('/merge', methods=['POST'])
def merge_pdfs():
    files = request.files.getlist('files')
    merged_pdf = fitz.open()
    for f in files:
        doc = fitz.open(stream=f.read(), filetype='pdf')
        merged_pdf.insert_pdf(doc)
    buffer = io.BytesIO()
    merged_pdf.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='unido.pdf')


@app.route('/split', methods=['POST'])
def split_pdfs():
    print(f"[{datetime.now()}] REQUISIÇÃO /split recebida.")

    if 'pdfs' not in request.files:
        return jsonify({"message": "Nenhum arquivo enviado."}), 400

    f = request.files.getlist('pdfs')[0]
    mode = request.form.get('mode')

    repair_needed = request.form.get('repair_pdf') == 'true'

    try:
        pdf_bytes = f.read()

        if repair_needed:
            print(f"[{datetime.now()}] Reparo solicitado. Limpando PDF em memória...")
            original_doc = fitz.open(stream=pdf_bytes, filetype='pdf')
            repaired_buffer = io.BytesIO()
            original_doc.save(repaired_buffer, garbage=4, deflate=True, clean=True)
            original_doc.close()
            pdf_bytes = repaired_buffer.getvalue()
            print(f"[{datetime.now()}] PDF reparado com sucesso.")

        pdf_doc = fitz.open(stream=pdf_bytes, filetype='pdf')
        print(f"[{datetime.now()}] PDF aberto. {pdf_doc.page_count} páginas.")
    except Exception as e:
        print(f"[{datetime.now()}] Erro ao abrir/reparar PDF: {e}")
        return jsonify({"message": f"Arquivo PDF inválido ou corrompido demais para reparar: {e}"}), 400

    filename = f.filename.rsplit('.', 1)[0]
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        if mode == 'parts':
            try:
                parts = int(request.form.get('parts'))
                if parts <= 0:
                    return jsonify({"message": "O número de partes deve ser maior que 0."}), 400
                if parts > pdf_doc.page_count:
                    return jsonify({"message": f"O número de partes ({parts}) não pode ser maior que o número de páginas ({pdf_doc.page_count})."}), 400
            except (ValueError, TypeError):
                return jsonify({"message": "Número de partes inválido."}), 400

            pages_per_part = math.ceil(pdf_doc.page_count / parts)
            for i in range(parts):
                start_page = i * pages_per_part
                end_page = min(start_page + pages_per_part - 1, pdf_doc.page_count - 1)
                if start_page > end_page:
                    continue
                new_pdf = fitz.open()
                for page_num in range(start_page, end_page + 1):
                    page = pdf_doc[page_num]
                    new_pdf.new_page(width=page.rect.width, height=page.rect.height)
                    new_pdf[-1].show_pdf_page(new_pdf[-1].rect, pdf_doc, page_num)
                output_buffer = io.BytesIO()
                new_pdf.save(output_buffer, garbage=4, deflate=True)
                output_buffer.seek(0)
                zipf.writestr(f"{filename}_parte_{i+1}_de_{parts}.pdf", output_buffer.read())
                new_pdf.close()

        elif mode == 'size':
            try:
                max_size_mb = float(request.form.get('max_size_mb'))
                if max_size_mb <= 0:
                    return jsonify({"message": "O tamanho máximo deve ser maior que 0 MB."}), 400
            except (ValueError, TypeError):
                return jsonify({"message": "Tamanho máximo inválido."}), 400

            max_size_bytes = max_size_mb * 1024 * 1024
            part_number = 1
            chunk_start_page = 0

            print(f"[{datetime.now()}] Iniciando divisão por tamanho (busca binária). Limite: {max_size_mb}MB.")

            while chunk_start_page < pdf_doc.page_count:
                print(f"[{datetime.now()}] ===== Bloco #{part_number} (pág. {chunk_start_page + 1}) =====")

                # checa página única
                single_page_doc = fitz.open()
                page = pdf_doc[chunk_start_page]
                single_page_doc.new_page(width=page.rect.width, height=page.rect.height)
                single_page_doc[-1].show_pdf_page(single_page_doc[-1].rect, pdf_doc, chunk_start_page)
                single_page_buffer = io.BytesIO()
                single_page_doc.save(single_page_buffer)
                if single_page_buffer.tell() > max_size_bytes:
                    return jsonify({"message": f"A página {chunk_start_page + 1} sozinha ({single_page_buffer.tell() / (1024*1024):.2f}MB) já é maior que o limite de {max_size_mb} MB."}), 400

                low, high = chunk_start_page, pdf_doc.page_count - 1
                best_end_page = chunk_start_page

                while low <= high:
                    mid = (low + high) // 2
                    test_doc = fitz.open()
                    for page_num in range(chunk_start_page, mid + 1):
                        p = pdf_doc[page_num]
                        test_doc.new_page(width=p.rect.width, height=p.rect.height)
                        test_doc[-1].show_pdf_page(test_doc[-1].rect, pdf_doc, page_num)
                    test_buffer = io.BytesIO()
                    test_doc.save(test_buffer)
                    test_doc.close()

                    if test_buffer.tell() <= max_size_bytes:
                        best_end_page = mid
                        low = mid + 1
                    else:
                        high = mid - 1

                final_chunk_doc = fitz.open()
                for page_num in range(chunk_start_page, best_end_page + 1):
                    p = pdf_doc[page_num]
                    final_chunk_doc.new_page(width=p.rect.width, height=p.rect.height)
                    final_chunk_doc[-1].show_pdf_page(final_chunk_doc[-1].rect, pdf_doc, page_num)
                output_buffer = io.BytesIO()
                final_chunk_doc.save(output_buffer, garbage=4, deflate=True)
                output_buffer.seek(0)
                zipf.writestr(f"{filename}_parte_{part_number}.pdf", output_buffer.read())
                final_chunk_doc.close()

                part_number += 1
                chunk_start_page = best_end_page + 1

    pdf_doc.close()
    zip_buffer.seek(0)
    print(f"[{datetime.now()}] Finalizado. Enviando ZIP.")
    return send_file(zip_buffer, as_attachment=True, download_name=f'{filename}_dividido.zip', mimetype='application/zip')


@app.route('/organize', methods=['POST'])
def organize_pdf():
    file = request.files.get('pdf')
    new_order = json.loads(request.form.get('order'))
    pdf_doc = fitz.open(stream=file.read(), filetype='pdf')
    new_doc = fitz.open()
    for item in new_order:
        # Rotação é aplicada na cópia (insert_pdf mantém a geometria)
        page_index = item['page'] - 1
        new_doc.insert_pdf(pdf_doc, from_page=page_index, to_page=page_index, rotate=item.get('rotation', 0))
    buffer = io.BytesIO()
    new_doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='organized.pdf')


# ---------------------------- No Cache ----------------------------
@app.after_request
def no_cache(response):
    response.headers['Cache-Control'] = 'no-store'
    return response


if __name__ == '__main__':
    # Para Linux
    app.run(debug=True, host='0.0.0.0', port=5009)
